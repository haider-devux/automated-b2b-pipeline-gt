"""
Generate a supervisor-facing Word document (.docx) listing every feature + safeguard of the pipeline,
as concise bullet points. Run:  python scripts/make_features_doc.py   ->  Granjur_Pipeline_Features_updated.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent  # scripts/ -> project root

SECTIONS = [
    ("1. Lead Discovery — finding companies (all free sources)", [
        "OpenStreetMap (Overpass API): local businesses worldwide, by city + category",
        "Free job boards (RemoteOK, Remotive): companies actively hiring developers — high buying intent",
        "CSV import: bring your own contact lists",
        "Bulk public registries (UK Companies House, People Data Labs, city business-license exports)",
        "Automatic de-duplication: the same company is never contacted twice",
        "Region targeting: US, EU, UK, GCC, CN, AU",
        "Big-brand blocklist: skips national chains and enterprises, focuses on small businesses",
    ]),
    ("2. Enrichment — getting the contact + context", [
        "Finds a real email from the company's own website (homepage + Contact / About / Team pages)",
        "Reads mailto: links first for accuracy (avoids scraping errors)",
        "Detects the website's tech stack (Shopify, WordPress, React, etc.)",
        "Optional website-speed score via Google PageSpeed",
        "Validates every email with a live DNS / MX check",
        "Fault-tolerant: one failing lookup never stops the whole batch",
    ]),
    ("3. Qualification & AI Pitch Writing", [
        "Scores and segments each lead (Legacy business / Funded startup / Low-tech e-commerce)",
        "Auto-disqualifies out-of-fit leads (e.g. too large, internal IT team)",
        "Writes a unique, plain-English pitch per lead with a local AI model",
        "Grounds every pitch in real facts from the company's site (anti-hallucination guard)",
        "Adds the recipient's first name when known",
        "English-only, non-technical, non-spammy tone",
    ]),
    ("4. Outreach & Sending", [
        "Sends via Gmail with a branded HTML signature (logo + 3 office addresses)",
        "“Book a 15-min call” button linked to Google Calendar (shows only real free slots)",
        "Human-approval mode OR fully automatic approval",
        "Test mode: previews every email to your own inbox first — nothing reaches companies",
        "Dry-run by default: nothing is sent until you deliberately switch to live",
    ]),
    ("5. Follow-ups (automatic drip)", [
        "4-step non-spammy sequence: nudge → issue → check-in → break-up",
        "Sent on a spaced cadence (day 3, 10, 20, 34 after the first email)",
        "Auto-stops the instant a prospect replies or books a call",
        "No cliché “just bumping this” lines; low-friction, interest-based CTAs",
    ]),
    ("6. Deliverability & Anti-Spam Safeguards (the checks)", [
        "Warmup cap: ramps daily volume for a fresh mailbox to protect sender reputation",
        "Send-window gate: never emails at 2 AM local time, on weekends, or on public holidays",
        "Role-inbox filter: skips info@ / sales@ for enterprises (spam traps); allows them only for small local businesses",
        "Global suppression list: bounced / unsubscribed / complained addresses are never mailed again",
        "Bounce scanning: hard bounces are detected and auto-suppressed",
        "One-click unsubscribe + real physical address in every email (legal compliance)",
        "Pre-send re-validation: dead addresses are dropped right before sending",
        "Consent-region guard: opt-in-only countries (DE / AT) are routed to manual review",
        "SPF / DKIM / DMARC monitoring on the Health page",
        "MX check before any guessed email address (no blind guessing that would bounce)",
    ]),
    ("7. Daily Quota System", [
        "Fixed target of 19 send-ready leads per day (configurable)",
        "Only genuinely sendable leads count — errors and “needs-contact” are excluded",
        "Quota-filler tops up any shortfall from OpenStreetMap + your fallback contact list",
        "Never overshoots the daily target or the mailbox warmup cap",
    ]),
    ("8. Automation & Reliability", [
        "One single command runs the entire pipeline end-to-end — no manual steps",
        "Region isolation and error-resilience (keeps going through individual failures)",
        "Resumable: safe to stop and restart without losing or double-sending leads",
    ]),
    ("9. Dashboard (web interface)", [
        "Live pipeline funnel: lead counts at every stage",
        "Leads hub with instant tag filters",
        "Outreach review, per-region send board, Follow-ups, Health, and in-depth Analytics pages",
    ]),
    ("10. Analytics & Per-Email Tracking", [
        "Per-email table: every sent email with delivered / open / click / reply / booked status",
        "Click any email to open its full timeline (discovered → sent → opened → clicked → replied)",
        "Open & click tracking via a 1x1 pixel + wrapped links — free, no paid analytics provider",
        "Clicks broken down by link type: website vs calendar/booking vs unsubscribe",
        "Conversion cohorts: by discovery source, and personal vs role (info@) inbox",
        "Daily sends vs the 19/day quota — see whether the target was met each day",
        "Conversion by segment and by region",
    ]),
    ("11. Reporting", [
        "Everything is Excel (.xlsx) — no CSV clutter; opens natively in Excel 2010 and newer",
        "One central database file (granjur_central.xlsx) accumulates every run's info together",
        "Central file sheets: Summary (current counts), Runs Log (one row per run over time), Latest Leads (full live list)",
        "Per-run snapshots (granjur_report_<date>.xlsx) freeze each run; the folder auto-prunes to the latest few",
        "One-click 'Download central Excel database' regenerates live from the database every time",
    ]),
    ("12. Cost & Privacy", [
        "100% free stack: Gmail + local AI + PostgreSQL + OpenStreetMap — no paid APIs",
        "The AI runs locally — company data never leaves the machine",
    ]),
]


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Granjur B2B Cold-Outreach Pipeline", level=0)
    sub = doc.add_paragraph("Features & Safeguards — Overview for Review")
    sub.runs[0].italic = True
    intro = doc.add_paragraph(
        "A free, self-hosted system that discovers businesses, writes personalized cold emails with AI, "
        "sends them safely from a single Gmail account, and follows up — fully automated, with strict "
        "anti-spam safeguards at every step.")
    intro.runs[0].font.size = Pt(11)

    for heading, bullets in SECTIONS:
        h = doc.add_heading(heading, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Note: the system runs in safe dry-run / test mode by default; live sending is a deliberate, "
        "separate switch. All sending stays within the mailbox warmup cap to protect deliverability.")
    foot.runs[0].italic = True
    foot.runs[0].font.size = Pt(9)

    # Canonical, single doc kept at the project root (see the folder-cleanup rule in Guide.md).
    out = ROOT / "Granjur_Pipeline_Features_updated.docx"
    try:
        doc.save(out)
    except PermissionError:                 # file open in Word -> write a side copy instead of failing
        out = ROOT / "Granjur_Pipeline_Features_updated_new.docx"
        doc.save(out)
        print("(the doc was open in Word — saved to _new; close Word and rename over the original)")
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
